"""Exportación de la transcripción como documento de trabajo editorial.

Distinto del export de corpus (export_service.py → JSON para máquinas) y del
bundle OKF (para otros agentes de IA): esto es un texto para que lo lea y lo
corrija una persona, con la atribución de cada intervención y su minuto.

Tres formatos, tres usos reales:
- Word (.docx): el formato en el que se corrige y se entrega a una revista.
- Markdown (.md): control de versiones y trasvase a otras herramientas.
- SubRip (.srt): subtitular el video con lo ya corregido.

Los tres llevan la misma ficha de procedencia (de dónde salió el material,
con qué modelo se transcribió) y la misma advertencia de que es una
transcripción automática pendiente de cotejo con el audio. Eso no es
decoración: publicar una transcripción exige poder citar la fuente y no
confundir el borrador de máquina con el texto revisado.
"""

from __future__ import annotations

import sqlite3
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

from videoindex.domain.diarization import (
    Intervencion,
    agrupar_intervenciones,
    nombre_visible,
)
from videoindex.infrastructure.db.repositories import (
    ProjectRepo,
    SegmentRepo,
    SpeakerRepo,
    VideoRepo,
)

ADVERTENCIA = (
    "Transcripción automática (voz a texto) sin revisar. Las etiquetas de "
    "hablante las asigna un modelo y pueden confundirse en cruces de voz. "
    "Antes de publicar: cotejar con el audio las citas textuales, los "
    "nombres propios y las cifras."
)


@dataclass
class TranscripcionExportable:
    """Todo lo que necesitan los tres formatos, resuelto una sola vez."""

    titulo: str
    intervenciones: list[Intervencion]
    nombres: dict[str, str]
    ficha: dict[str, str]

    @property
    def hablantes(self) -> list[str]:
        vistos: list[str] = []
        for i in self.intervenciones:
            nombre = nombre_visible(i.speaker, self.nombres)
            if nombre not in vistos:
                vistos.append(nombre)
        return vistos


def marca_tiempo(segundos: float, con_milisegundos: bool = False) -> str:
    """hh:mm:ss (o hh:mm:ss,mmm en SRT, que exige coma decimal)."""
    total = max(0.0, segundos)
    h, resto = divmod(int(total), 3600)
    m, s = divmod(resto, 60)
    if con_milisegundos:
        ms = int(round((total - int(total)) * 1000))
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"
    return f"{h:02d}:{m:02d}:{s:02d}"


def preparar(
    con: sqlite3.Connection, video_id: str, pausa_maxima_s: float = 0.0
) -> TranscripcionExportable:
    """Reúne transcripción + nombres de hablante + ficha de procedencia.

    Lanza ValueError si el video no existe o no tiene transcripción: exportar
    un documento vacío sería peor que avisar.
    """
    video = VideoRepo(con).por_id(video_id)
    if video is None:
        raise ValueError(f"Video no encontrado: {video_id}")
    segmentos = SegmentRepo(con).por_video(video_id)
    if not segmentos:
        raise ValueError(
            f'"{video.title}" todavía no tiene transcripción: procésalo antes de exportar.'
        )

    ficha = {"Título": video.title}
    if video.source_url:
        ficha["Fuente"] = video.source_url
    if video.source_channel:
        ficha["Canal / autor"] = video.source_channel
    if video.source_published_at:
        ficha["Publicado"] = video.source_published_at
    if video.duration_seconds:
        ficha["Duración"] = marca_tiempo(video.duration_seconds)
    proyecto = {p.project_id: p.name for p in ProjectRepo(con).listar()}.get(video.project_id)
    if proyecto:
        ficha["Proyecto"] = proyecto
    ficha["Archivo"] = Path(video.path).name
    ficha["Exportado"] = datetime.now(UTC).astimezone().strftime("%Y-%m-%d %H:%M")

    return TranscripcionExportable(
        titulo=video.title,
        intervenciones=agrupar_intervenciones(segmentos, pausa_maxima_s),
        nombres=SpeakerRepo(con).nombres(video_id),
        ficha=ficha,
    )


def a_markdown(t: TranscripcionExportable, con_timestamps: bool = True) -> str:
    lineas = [f"# {t.titulo}", ""]
    for clave, valor in t.ficha.items():
        lineas.append(f"- **{clave}:** {valor}")
    lineas += ["", f"> {ADVERTENCIA}", ""]
    if t.hablantes:
        lineas += [f"**Intervienen:** {', '.join(t.hablantes)}", ""]
    lineas.append("---")
    lineas.append("")
    for intervencion in t.intervenciones:
        quien = nombre_visible(intervencion.speaker, t.nombres)
        marca = f"`[{marca_tiempo(intervencion.start_time)}]` " if con_timestamps else ""
        lineas += [f"{marca}**{quien}:** {intervencion.texto}", ""]
    return "\n".join(lineas)


def a_srt(t: TranscripcionExportable) -> str:
    """Un subtítulo por intervención, con el nombre del hablante delante.

    Nota: las intervenciones pueden ser largas para un subtítulo; el archivo
    sirve para llevar el texto corregido a un editor de video, no para
    subtitular sin retocar."""
    bloques = []
    for numero, intervencion in enumerate(t.intervenciones, 1):
        quien = nombre_visible(intervencion.speaker, t.nombres)
        bloques.append(
            f"{numero}\n"
            f"{marca_tiempo(intervencion.start_time, True)} --> "
            f"{marca_tiempo(intervencion.end_time, True)}\n"
            f"{quien}: {intervencion.texto}\n"
        )
    return "\n".join(bloques)


def exportar_markdown(
    con: sqlite3.Connection, video_id: str, destino: str | Path, con_timestamps: bool = True
) -> Path:
    destino = Path(destino)
    destino.write_text(a_markdown(preparar(con, video_id), con_timestamps), encoding="utf-8")
    return destino


def exportar_srt(con: sqlite3.Connection, video_id: str, destino: str | Path) -> Path:
    destino = Path(destino)
    destino.write_text(a_srt(preparar(con, video_id)), encoding="utf-8")
    return destino


def exportar_docx(
    con: sqlite3.Connection, video_id: str, destino: str | Path, con_timestamps: bool = True
) -> Path:
    """Word con estilos reales (Título, Subtítulo, Cita), no texto plano
    metido en un .docx: así el documento entra directo al flujo de edición
    y los estilos se pueden reemplazar por los de la revista."""
    from docx import Document
    from docx.shared import Pt

    t = preparar(con, video_id)
    destino = Path(destino)

    doc = Document()
    doc.add_heading(t.titulo, level=0)

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Light Grid Accent 1"
    for clave, valor in t.ficha.items():
        fila = tabla.add_row().cells
        fila[0].text = clave
        fila[1].text = str(valor)

    doc.add_paragraph()
    aviso = doc.add_paragraph(ADVERTENCIA)
    aviso.style = doc.styles["Intense Quote"]

    if t.hablantes:
        doc.add_heading("Intervienen", level=2)
        for nombre in t.hablantes:
            doc.add_paragraph(nombre, style="List Bullet")

    doc.add_heading("Transcripción", level=1)
    for intervencion in t.intervenciones:
        quien = nombre_visible(intervencion.speaker, t.nombres)
        parrafo = doc.add_paragraph()
        if con_timestamps:
            marca = parrafo.add_run(f"[{marca_tiempo(intervencion.start_time)}] ")
            marca.font.size = Pt(8)
            marca.italic = True
        etiqueta = parrafo.add_run(f"{quien}: ")
        etiqueta.bold = True
        parrafo.add_run(intervencion.texto)

    doc.save(str(destino))
    return destino


FORMATOS = {
    "docx": ("Word (.docx)", exportar_docx),
    "md": ("Markdown (.md)", exportar_markdown),
    "srt": ("Subtítulos (.srt)", exportar_srt),
}
