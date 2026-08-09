"""Paquete de entrega de una transcripción profesional.

Produce el juego completo de documentos con los que se trabaja una
transcripción destinada a publicarse:

  transcripcion_literal.docx     lo que se dijo, tal cual
  transcripcion_limpia.docx      legible, sin alterar el contenido
  transcripcion_completa.txt     la limpia en texto plano
  subtitulos.srt                 para el video
  participantes_identificados.xlsx   quién es quién, con su evidencia
  citas_literarias.xlsx          los textos leídos en voz alta
  incertidumbres.md              lo que tiene que resolver una persona
  proceso_tecnico.md             cómo se produjo todo esto

Dos principios recorren el módulo:

1. **Nunca inventar.** Un nombre solo aparece si hay evidencia en el propio
   video (un rótulo, una presentación en voz alta). Si no la hay, se usa una
   identificación funcional y el caso se anota en `incertidumbres.md`.
2. **Separar lo automático de lo humano.** Los documentos dicen en su primera
   página que son transcripción de máquina sin cotejar, y `incertidumbres.md`
   concentra lo que hay que revisar en vez de repartirlo por el texto.
"""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from videoindex.application.identificacion_service import CitaEnPantalla, Identidad
from videoindex.application.transcript_export_service import ADVERTENCIA, marca_tiempo
from videoindex.domain.alucinaciones import es_alucinacion_probable
from videoindex.domain.diarization import Intervencion, agrupar_intervenciones
from videoindex.domain.glosario import corregir
from videoindex.domain.limpieza import limpiar_para_lectura
from videoindex.domain.models import TranscriptSegment

# Confianza media de Whisper por debajo de la cual conviene que alguien
# escuche el pasaje. exp(avg_logprob): 0.6 ya es una zona de dudas.
CONFIANZA_DUDOSA = 0.60
# Solo se listan los pasajes dudosos más graves; una lista de doscientos
# avisos no la revisa nadie.
MAX_PASAJES_DUDOSOS = 25


@dataclass
class Contexto:
    """De dónde salió el material y con qué se procesó."""

    titulo: str
    archivo: str
    duracion_s: float
    url: str | None = None
    canal: str | None = None
    publicado: str | None = None
    modelo_transcripcion: str = ""
    modelo_diarizacion: str = ""
    modelo_ocr: str = ""
    notas: list[str] = field(default_factory=list)


def _nombre_de(intervencion: Intervencion, identidades: dict[str, Identidad]) -> Identidad | None:
    return identidades.get(intervencion.speaker) if intervencion.speaker else None


def _etiqueta(
    intervencion: Intervencion, identidades: dict[str, Identidad], ya_presentados: set[str]
) -> str:
    """Nombre + cargo la PRIMERA vez; solo el nombre después.

    Repetir el cargo en cada intervención satura la lectura; omitirlo desde
    el principio deja al lector sin saber quién es cada quien.
    """
    identidad = _nombre_de(intervencion, identidades)
    if identidad is None:
        return "VOZ NO IDENTIFICADA"
    if identidad.speaker_label in ya_presentados:
        return identidad.etiqueta_editorial
    ya_presentados.add(identidad.speaker_label)
    return identidad.descriptor


def _es_lectura(
    intervencion: Intervencion, citas: Sequence[CitaEnPantalla]
) -> CitaEnPantalla | None:
    """La cita cuya tarjeta aparece durante esta intervención, si la hay."""
    for cita in citas:
        if intervencion.start_time - 4.0 <= cita.inicio_s <= intervencion.end_time:
            return cita
    return None


# Un párrafo de más de minuto y medio de habla (~250 palabras) ya es
# incómodo de corregir en Word. Se corta ahí, siempre en frontera de
# segmento y sin cambiar de hablante.
DURACION_PARRAFO_S = 90.0
PAUSA_PARRAFO_S = 8.0


def construir_intervenciones(
    segmentos: list[TranscriptSegment], fin_contenido_s: float | None = None
) -> list[Intervencion]:
    """Intervenciones legibles: además del cambio de hablante, se parten por
    silencios y por duración, para que la narración continua de un documental
    no salga como un solo párrafo de veinte minutos.

    `fin_contenido_s` deja fuera lo que suene ya sobre los créditos finales
    (música, canción, una última locución de cierre): no es contenido del
    documental y no debe entrar en el texto que se publica.
    """
    utiles = (
        [s for s in segmentos if s.start_time < fin_contenido_s]
        if fin_contenido_s is not None
        else list(segmentos)
    )
    intervenciones = agrupar_intervenciones(
        utiles, pausa_maxima_s=PAUSA_PARRAFO_S, duracion_maxima_s=DURACION_PARRAFO_S
    )
    # Se filtra sobre la intervención ya montada, no sobre el segmento suelto:
    # la coletilla suele venir troceada y solo se reconoce entera.
    return [i for i in intervenciones if not es_alucinacion_probable(i.texto)]


def alucinaciones_descartadas(
    segmentos: list[TranscriptSegment], fin_contenido_s: float | None = None
) -> list[Intervencion]:
    """Las intervenciones que se dejaron fuera por parecer inventadas por el
    modelo. Van a `incertidumbres.md`: nada se descarta en silencio."""
    utiles = (
        [s for s in segmentos if s.start_time < fin_contenido_s]
        if fin_contenido_s is not None
        else list(segmentos)
    )
    todas = agrupar_intervenciones(
        utiles, pausa_maxima_s=PAUSA_PARRAFO_S, duracion_maxima_s=DURACION_PARRAFO_S
    )
    return [i for i in todas if es_alucinacion_probable(i.texto)]


# --------------------------------------------------------------------------
# Documentos
# --------------------------------------------------------------------------


def _nuevo_docx(contexto: Contexto, subtitulo: str, identidades: list[Identidad]):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(contexto.titulo, level=0)
    doc.add_paragraph(subtitulo, style="Subtitle")

    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Light Grid Accent 1"
    ficha = {
        "Archivo": contexto.archivo,
        "Duración": marca_tiempo(contexto.duracion_s),
        "Fuente": contexto.url or "—",
        "Canal / autor": contexto.canal or "—",
        "Publicado": contexto.publicado or "—",
        "Transcripción": contexto.modelo_transcripcion,
        "Separación de voces": contexto.modelo_diarizacion,
        "Generado": datetime.now().astimezone().strftime("%Y-%m-%d %H:%M"),
    }
    for clave, valor in ficha.items():
        celdas = tabla.add_row().cells
        celdas[0].text = clave
        celdas[1].text = str(valor)

    doc.add_paragraph()
    doc.add_paragraph(ADVERTENCIA, style="Intense Quote")

    doc.add_heading("Participantes", level=1)
    tabla_p = doc.add_table(rows=1, cols=4)
    tabla_p.style = "Light Grid Accent 1"
    for celda, titulo in zip(
        tabla_p.rows[0].cells, ("Identificación", "Función", "Confianza", "Aparece"), strict=True
    ):
        celda.text = titulo
        celda.paragraphs[0].runs[0].bold = True
    for identidad in identidades:
        celdas = tabla_p.add_row().cells
        celdas[0].text = identidad.etiqueta_editorial
        celdas[1].text = ", ".join(x for x in (identidad.funcion, identidad.institucion) if x) or (
            "NARRACIÓN" if identidad.es_voz_en_off else "—"
        )
        celdas[2].text = identidad.confianza
        celdas[3].text = marca_tiempo(identidad.primera_aparicion)

    doc.add_page_break()
    estilo = doc.styles["Normal"]
    estilo.font.size = Pt(11)
    return doc


def texto_de_lectura(texto: str, glosario: Sequence[str] | None) -> tuple[str, list]:
    """Versión legible con la ortografía de los nombres corregida.

    El glosario solo se aplica a la versión LIMPIA. La literal es el registro
    auditable de lo que produjo la máquina: si se corrigiera también ahí, no
    quedaría contra qué cotejar.
    """
    legible = limpiar_para_lectura(texto)
    if not glosario:
        return legible, []
    resultado = corregir(legible, list(glosario))
    return resultado.texto, resultado.sugerencias


def _escribir_cuerpo(doc, intervenciones, identidades, citas, limpiar: bool, glosario=None) -> None:
    from docx.shared import Pt, RGBColor

    doc.add_heading("Transcripción", level=1)
    ya_presentados: set[str] = set()
    cita_abierta: CitaEnPantalla | None = None

    for intervencion in intervenciones:
        cita = _es_lectura(intervencion, citas)
        if cita is not None and cita is not cita_abierta:
            cita_abierta = cita
            aviso = doc.add_paragraph()
            marca = aviso.add_run(
                f"[{marca_tiempo(intervencion.start_time)}] INICIO — LECTURA DE TEXTO"
            )
            marca.bold = True
            detalle = doc.add_paragraph(
                f"Título en pantalla: «{cita.titulo}»"
                + (f" · Autor: {cita.autor}" if cita.autor else "")
                + (f" · {cita.anio}" if cita.anio else "")
            )
            detalle.runs[0].italic = True

        texto = texto_de_lectura(intervencion.texto, glosario)[0] if limpiar else intervencion.texto
        parrafo = doc.add_paragraph()
        sello = parrafo.add_run(f"[{marca_tiempo(intervencion.start_time)}] ")
        sello.font.size = Pt(8)
        sello.italic = True
        sello.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        quien = parrafo.add_run(f"{_etiqueta(intervencion, identidades, ya_presentados)}\n")
        quien.bold = True
        parrafo.add_run(texto)

        if cita_abierta is not None and cita is None:
            fin = doc.add_paragraph()
            fin.add_run(
                f"[{marca_tiempo(intervencion.start_time)}] FIN — LECTURA DE TEXTO"
            ).bold = True
            cita_abierta = None


def escribir_docx(
    destino: Path,
    contexto: Contexto,
    intervenciones: list[Intervencion],
    identidades: list[Identidad],
    citas: Sequence[CitaEnPantalla],
    limpiar: bool,
    glosario: Sequence[str] | None = None,
) -> Path:
    subtitulo = (
        "Transcripción editorial (versión de lectura)"
        if limpiar
        else "Transcripción literal (sin editar)"
    )
    doc = _nuevo_docx(contexto, subtitulo, identidades)
    _escribir_cuerpo(doc, intervenciones, {i.speaker_label: i for i in identidades}, citas, limpiar)
    doc.save(str(destino))
    return destino


def escribir_txt(
    destino: Path,
    contexto: Contexto,
    intervenciones: list[Intervencion],
    identidades: list[Identidad],
    glosario: Sequence[str] | None = None,
) -> Path:
    indice = {i.speaker_label: i for i in identidades}
    ya: set[str] = set()
    lineas = [contexto.titulo, "=" * len(contexto.titulo), ""]
    if contexto.url:
        lineas.append(f"Fuente: {contexto.url}")
    lineas += [f"Duración: {marca_tiempo(contexto.duracion_s)}", "", ADVERTENCIA, "", "-" * 70, ""]
    for intervencion in intervenciones:
        lineas.append(
            f"[{marca_tiempo(intervencion.start_time)}] {_etiqueta(intervencion, indice, ya)}"
        )
        lineas += [texto_de_lectura(intervencion.texto, glosario)[0], ""]
    destino.write_text("\n".join(lineas), encoding="utf-8")
    return destino


def escribir_srt(
    destino: Path, intervenciones: list[Intervencion], identidades: list[Identidad]
) -> Path:
    indice = {i.speaker_label: i for i in identidades}
    bloques = []
    for numero, intervencion in enumerate(intervenciones, 1):
        identidad = indice.get(intervencion.speaker or "")
        quien = identidad.etiqueta_editorial if identidad else "VOZ NO IDENTIFICADA"
        bloques.append(
            f"{numero}\n"
            f"{marca_tiempo(intervencion.start_time, True)} --> "
            f"{marca_tiempo(intervencion.end_time, True)}\n"
            f"{quien}: {limpiar_para_lectura(intervencion.texto)}\n"
        )
    destino.write_text("\n".join(bloques), encoding="utf-8")
    return destino


def escribir_participantes(destino: Path, identidades: list[Identidad]) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    libro = Workbook()
    hoja = libro.active
    hoja.title = "Participantes"
    encabezados = [
        "Identificación",
        "Nombre",
        "Función",
        "Institución",
        "ID de voz",
        "Primera aparición",
        "Última aparición",
        "Minutos hablados",
        "Nivel de confianza",
        "Evidencia",
        "Observaciones",
    ]
    hoja.append(encabezados)
    for celda in hoja[1]:
        celda.font = Font(bold=True)

    for identidad in identidades:
        hoja.append(
            [
                identidad.etiqueta_editorial,
                identidad.nombre or "",
                identidad.funcion or ("NARRACIÓN" if identidad.es_voz_en_off else ""),
                identidad.institucion or "",
                identidad.speaker_label,
                marca_tiempo(identidad.primera_aparicion),
                marca_tiempo(identidad.ultima_aparicion),
                round(identidad.segundos_hablados / 60, 1),
                identidad.confianza,
                " | ".join(identidad.evidencias) or "Sin evidencia en el video",
                "Voz en off / narración" if identidad.es_voz_en_off else "",
            ]
        )
    for columna, ancho in zip(
        "ABCDEFGHIJK", (28, 24, 26, 14, 12, 16, 16, 16, 14, 70, 24), strict=True
    ):
        hoja.column_dimensions[columna].width = ancho
    libro.save(str(destino))
    return destino


def escribir_citas(destino: Path, citas: Sequence[CitaEnPantalla]) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font

    libro = Workbook()
    hoja = libro.active
    hoja.title = "Citas literarias"
    hoja.append(
        [
            "Nº",
            "Inicio",
            "Fin",
            "Título (según pantalla)",
            "Autor",
            "Año",
            "Evidencia",
            "Nivel de confianza",
            "Observaciones",
        ]
    )
    for celda in hoja[1]:
        celda.font = Font(bold=True)
    for numero, cita in enumerate(citas, 1):
        hoja.append(
            [
                numero,
                marca_tiempo(cita.inicio_s),
                marca_tiempo(cita.fin_s),
                cita.titulo,
                cita.autor or "No identificado",
                cita.anio or "",
                f"Tarjeta sobreimpresa en {marca_tiempo(cita.inicio_s)}",
                "ALTO" if cita.autor and cita.anio else "MEDIO",
                "El inicio y el fin del pasaje leído requieren cotejo con el audio",
            ]
        )
    for columna, ancho in zip("ABCDEFGHI", (6, 12, 12, 42, 26, 8, 40, 16, 52), strict=True):
        hoja.column_dimensions[columna].width = ancho
    libro.save(str(destino))
    return destino


def escribir_incertidumbres(
    destino: Path,
    identidades: list[Identidad],
    citas: Sequence[CitaEnPantalla],
    segmentos: list[TranscriptSegment],
    descartadas: Sequence[Intervencion] = (),
    sugerencias_glosario: Sequence = (),
) -> tuple[Path, int]:
    """Solo lo que necesita decisión humana. Devuelve también cuántos asuntos
    hay: un documento de incertidumbres que nadie cuenta no sirve de nada."""
    lineas = [
        "# Asuntos que requieren revisión humana",
        "",
        "Cada punto es algo que el procesamiento automático **no** pudo resolver "
        "con evidencia suficiente. No se han inventado nombres ni atribuciones.",
        "",
    ]
    asuntos = 0

    sin_nombre = [i for i in identidades if not i.nombre]
    if sin_nombre:
        lineas += ["## Voces sin identificar", ""]
        for identidad in sin_nombre:
            asuntos += 1
            papel = "voz en off / narración" if identidad.es_voz_en_off else "sin función clara"
            lineas += [
                f"### {marca_tiempo(identidad.primera_aparicion)} — {identidad.speaker_label}",
                "",
                f"Habla {identidad.segundos_hablados / 60:.1f} min en total "
                f"(hasta {marca_tiempo(identidad.ultima_aparicion)}). Tratada como {papel}.",
                "",
                "No aparece ningún rótulo sobre sus intervenciones ni se la presenta "
                "en voz alta. Si su nombre figura en los créditos finales, conviene "
                "asignarlo a mano.",
                "",
            ]

    dudosas = [i for i in identidades if i.nombre and i.confianza != "ALTO"]
    if dudosas:
        lineas += ["## Identificaciones que no llegan a confianza alta", ""]
        for identidad in dudosas:
            asuntos += 1
            lineas += [
                f"### {marca_tiempo(identidad.primera_aparicion)} — «{identidad.nombre}» "
                f"({identidad.confianza})",
                "",
                *[f"- {e}" for e in identidad.evidencias],
                "",
            ]

    if citas:
        lineas += ["## Textos leídos en voz alta", ""]
        for cita in citas:
            asuntos += 1
            lineas += [
                f"### {marca_tiempo(cita.inicio_s)} — «{cita.titulo}»",
                "",
                f"Autor según la tarjeta: {cita.autor or 'no identificado'}"
                + (f" ({cita.anio})" if cita.anio else "")
                + ".",
                "",
                "La tarjeta da el título y el autor, pero **no dónde empieza y "
                "termina exactamente la lectura**: eso hay que fijarlo de oído. "
                "Tampoco se ha reconstruido la disposición en versos, que el "
                "reconocimiento de voz no puede deducir de la entonación.",
                "",
            ]

    if descartadas:
        lineas += [
            "## Pasajes omitidos por parecer inventados por el modelo",
            "",
            "Whisper rellena los tramos sin habla (música, silencio, créditos) con "
            "coletillas de subtítulos de YouTube o repitiendo una frase en bucle. "
            "Estos pasajes **no están** en la transcripción; se listan aquí para que "
            "puedas comprobarlo en el audio y recuperarlos si alguno era real.",
            "",
        ]
        for intervencion in descartadas:
            asuntos += 1
            lineas.append(
                f"- **{marca_tiempo(intervencion.start_time)}**: «{intervencion.texto[:160]}»"
            )
        lineas.append("")

    if sugerencias_glosario:
        vistas = {(s.original, s.corregido) for s in sugerencias_glosario}
        lineas += [
            "## Nombres que quizá haya que corregir",
            "",
            "Se parecen a un nombre que aparece escrito en los rótulos del video, "
            "pero no lo bastante como para cambiarlos sin mirar: **no se han "
            "aplicado**. Compruébalos y corrígelos a mano si procede.",
            "",
        ]
        for original, corregido in sorted(vistas):
            asuntos += 1
            lineas.append(f"- «{original}» → ¿«{corregido}»?")
        lineas.append("")

    peores = sorted(
        (s for s in segmentos if s.confidence < CONFIANZA_DUDOSA), key=lambda s: s.confidence
    )[:MAX_PASAJES_DUDOSOS]
    if peores:
        lineas += [
            "## Pasajes con reconocimiento dudoso",
            "",
            f"Los {len(peores)} fragmentos donde el modelo estuvo menos seguro "
            "(de peor a mejor). Suelen ser nombres propios, cifras y siglas: "
            "es donde conviene escuchar antes de publicar.",
            "",
        ]
        for seg in peores:
            asuntos += 1
            lineas.append(
                f"- **{marca_tiempo(seg.start_time)}** (confianza {seg.confidence:.0%}): "
                f"«{seg.clean_text.strip()}»"
            )
        lineas.append("")

    destino.write_text("\n".join(lineas), encoding="utf-8")
    return destino, asuntos


def escribir_proceso_tecnico(
    destino: Path,
    contexto: Contexto,
    identidades: list[Identidad],
    citas: Sequence[CitaEnPantalla],
    segmentos: list[TranscriptSegment],
    n_rotulos: int,
) -> Path:
    identificadas = [i for i in identidades if i.nombre]
    lineas = [
        "# Proceso técnico",
        "",
        "## Material",
        "",
        f"- Archivo: `{contexto.archivo}`",
        f"- Duración: {marca_tiempo(contexto.duracion_s)}",
        f"- Fuente: {contexto.url or 'no consignada'}",
        f"- Canal / autor: {contexto.canal or 'no consignado'}",
        f"- Publicado: {contexto.publicado or 'no consignado'}",
        "",
        "## Herramientas",
        "",
        f"- Transcripción: **{contexto.modelo_transcripcion}**, local, sin API.",
        f"- Separación de voces: **{contexto.modelo_diarizacion}**, local, sin API.",
        f"- Lectura de rótulos: **{contexto.modelo_ocr}** sobre fotogramas extraídos con PyAV.",
        "- Decodificación de audio y video: PyAV (no requiere ffmpeg instalado).",
        "- Documentos: python-docx y openpyxl.",
        "- Coste de API: **0 USD**. Todo el procesamiento corrió en esta máquina.",
        "",
        "## Resultados",
        "",
        f"- Segmentos de habla transcritos: {len(segmentos)}",
        f"- Voces distinguidas: {len(identidades)}",
        f"- Voces con nombre propio: {len(identificadas)}",
        f"- Voces sin identificar: {len(identidades) - len(identificadas)}",
        f"- Rótulos leídos en pantalla: {n_rotulos}",
        f"- Textos literarios anunciados por tarjeta: {len(citas)}",
        "",
        "## Cómo se identificó a cada persona",
        "",
        "El nombre NO sale del audio: sale de cruzar los turnos de voz con los "
        "rótulos sobreimpresos. Cada rótulo se lee muchas veces (un fotograma "
        "por segundo) y se consolida por consenso, porque el OCR de un "
        "fotograma suelto es poco fiable: el mismo rótulo puede leerse "
        "«SOLEDAD BIANCHI» en un cuadro y «%LEDAD BIANCHI» en el siguiente.",
        "",
        "Una identificación llega a confianza ALTA solo si el rótulo trae "
        "nombre y cargo, aparece mientras esa voz habla, y ningún otro rótulo "
        "reclama el mismo nombre para otra voz.",
        "",
        "## Límites conocidos",
        "",
        "- **Habla superpuesta**: cuando dos personas hablan a la vez, el tramo "
        "se atribuye a una sola voz. El método (embeddings de voz + "
        "agrupamiento) no separa voces simultáneas.",
        "- **Créditos finales**: no se han cruzado automáticamente con las voces "
        "sin identificar; es la vía más probable para resolverlas.",
        "- **Versos**: la disposición en versos de un texto leído no se puede "
        "deducir del audio. Los pasajes se marcan, no se versifican.",
        "- La transcripción **no ha sido cotejada con el audio por una persona**.",
        "",
    ]
    if contexto.notas:
        lineas += ["## Notas de esta ejecución", "", *[f"- {n}" for n in contexto.notas], ""]
    destino.write_text("\n".join(lineas), encoding="utf-8")
    return destino


def generar_paquete(
    carpeta: str | Path,
    contexto: Contexto,
    segmentos: list[TranscriptSegment],
    identidades: list[Identidad],
    citas: Sequence[CitaEnPantalla],
    n_rotulos: int = 0,
    fin_contenido_s: float | None = None,
    glosario: Sequence[str] | None = None,
) -> dict[str, Path]:
    """Escribe los ocho documentos y devuelve sus rutas.

    `fin_contenido_s`: instante en que arrancan los créditos finales. Lo que
    suene por encima de los créditos (música, una locución de cierre) no es
    contenido del documental y no entra en el texto.
    """
    carpeta = Path(carpeta)
    carpeta.mkdir(parents=True, exist_ok=True)
    intervenciones = construir_intervenciones(segmentos, fin_contenido_s)
    if fin_contenido_s is not None:
        contexto.notas.append(
            f"El texto termina en {marca_tiempo(fin_contenido_s)}: a partir de ahí "
            "corren los créditos finales, que se excluyen de la transcripción."
        )

    salidas: dict[str, Path] = {}
    salidas["literal"] = escribir_docx(
        carpeta / "transcripcion_literal.docx", contexto, intervenciones, identidades, citas, False
    )
    salidas["limpia"] = escribir_docx(
        carpeta / "transcripcion_limpia.docx",
        contexto,
        intervenciones,
        identidades,
        citas,
        True,
        glosario,
    )
    salidas["txt"] = escribir_txt(
        carpeta / "transcripcion_completa.txt", contexto, intervenciones, identidades, glosario
    )

    # Lo que el glosario corrigió y lo que solo sugiere, para el informe.
    aplicadas, sugeridas = [], []
    if glosario:
        for intervencion in intervenciones:
            resultado = corregir(limpiar_para_lectura(intervencion.texto), list(glosario))
            aplicadas += resultado.cambios
            sugeridas += resultado.sugerencias
        if aplicadas:
            contexto.notas.append(
                f"{len(aplicadas)} nombre(s) con la ortografía corregida contra el "
                "glosario sacado de los rótulos del video"
            )
    salidas["srt"] = escribir_srt(carpeta / "subtitulos.srt", intervenciones, identidades)
    salidas["participantes"] = escribir_participantes(
        carpeta / "participantes_identificados.xlsx", identidades
    )
    salidas["citas"] = escribir_citas(carpeta / "citas_literarias.xlsx", citas)
    descartadas = alucinaciones_descartadas(segmentos, fin_contenido_s)
    if descartadas:
        contexto.notas.append(
            f"{len(descartadas)} pasaje(s) omitido(s) por parecer alucinación del "
            "modelo; se listan en incertidumbres.md"
        )
    salidas["incertidumbres"], asuntos = escribir_incertidumbres(
        carpeta / "incertidumbres.md", identidades, citas, segmentos, descartadas, sugeridas
    )
    contexto.notas.append(f"{asuntos} asuntos consignados en incertidumbres.md")
    salidas["proceso"] = escribir_proceso_tecnico(
        carpeta / "proceso_tecnico.md", contexto, identidades, citas, segmentos, n_rotulos
    )
    return salidas
