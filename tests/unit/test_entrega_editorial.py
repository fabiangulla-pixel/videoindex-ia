"""Paquete de entrega: los ocho documentos de una transcripción profesional.

Cada documento se vuelve a ABRIR con su librería y se lee su contenido. Que
un archivo exista y pese unos kilobytes no prueba nada: lo que importa es que
diga quién habla, en qué minuto, y que no haya inventado ningún nombre.
"""

from __future__ import annotations

from uuid import uuid4

import pytest

from videoindex.application.entrega_editorial import (
    Contexto,
    construir_intervenciones,
    generar_paquete,
)
from videoindex.application.identificacion_service import ALTO, BAJO, CitaEnPantalla, Identidad
from videoindex.domain.models import TranscriptSegment


def _seg(inicio, fin, texto, speaker, confianza=0.9):
    return TranscriptSegment(
        segment_id=str(uuid4()),
        video_id="doc",
        start_time=inicio,
        end_time=fin,
        raw_text=texto,
        clean_text=texto,
        confidence=confianza,
        speaker=speaker,
    )


@pytest.fixture
def material():
    segmentos = [
        _seg(0.0, 40.0, "Eh, la la ciudad cambiaba en aquellos años.", "SPEAKER_00"),
        _seg(105.0, 130.0, "Confieso que he vivido.", "SPEAKER_00"),
        _seg(205.0, 240.0, "El exilio marcó a toda una generación.", "SPEAKER_01"),
        # Vuelve a hablar tras un silencio largo: es una intervención NUEVA, y
        # sirve para comprobar que el cargo ya no se repite en la segunda.
        _seg(300.0, 320.0, "Y eso se nota en la obra.", "SPEAKER_01", confianza=0.35),
        _seg(400.0, 430.0, "Nadie sabe quién soy.", "SPEAKER_02"),
    ]
    identidades = [
        Identidad(
            speaker_label="SPEAKER_00",
            confianza=BAJO,
            es_voz_en_off=True,
            primera_aparicion=0.0,
            ultima_aparicion=130.0,
            segundos_hablados=65.0,
            evidencias=["Habla 1 min y nunca aparece rotulada"],
        ),
        Identidad(
            speaker_label="SPEAKER_01",
            nombre="CARLA ULLOA",
            funcion="HISTORIADORA",
            confianza=ALTO,
            primera_aparicion=205.0,
            ultima_aparicion=320.0,
            segundos_hablados=55.0,
            evidencias=["Rótulo en pantalla [00:03:25]: «CARLA ULLOA / HISTORIADORA»"],
        ),
        Identidad(
            speaker_label="SPEAKER_02",
            confianza=BAJO,
            primera_aparicion=400.0,
            ultima_aparicion=430.0,
            segundos_hablados=30.0,
        ),
    ]
    citas = [
        CitaEnPantalla(
            inicio_s=104.0,
            fin_s=113.0,
            titulo="Confieso que he vivido",
            autor="Pablo Neruda",
            anio="1974",
        )
    ]
    contexto = Contexto(
        titulo="Documental de prueba",
        archivo="documental.m4a",
        duracion_s=3223.0,
        url="https://www.youtube.com/watch?v=demo",
        canal="Canal de prueba",
        publicado="2026-05-20",
        modelo_transcripcion="faster-whisper large-v3-turbo",
        modelo_diarizacion="ECAPA + agrupamiento",
        modelo_ocr="Tesseract 5",
    )
    return segmentos, identidades, citas, contexto


def test_se_generan_los_ocho_documentos(material, tmp_path):
    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas, n_rotulos=7)

    esperados = {
        "transcripcion_literal.docx",
        "transcripcion_limpia.docx",
        "transcripcion_completa.txt",
        "subtitulos.srt",
        "participantes_identificados.xlsx",
        "citas_literarias.xlsx",
        "incertidumbres.md",
        "proceso_tecnico.md",
    }
    assert {p.name for p in salidas.values()} == esperados
    assert all(p.exists() and p.stat().st_size > 0 for p in salidas.values())


def test_la_version_literal_conserva_las_vacilaciones_y_la_limpia_no(material, tmp_path):
    """Es la diferencia que justifica entregar dos documentos."""
    from docx import Document

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)

    literal = "\n".join(p.text for p in Document(str(salidas["literal"])).paragraphs)
    limpia = "\n".join(p.text for p in Document(str(salidas["limpia"])).paragraphs)

    assert "Eh, la la ciudad" in literal
    assert "Eh, la la ciudad" not in limpia
    assert "La ciudad cambiaba" in limpia  # el contenido sigue ahí


def test_el_cargo_solo_aparece_en_la_primera_intervencion(material, tmp_path):
    from docx import Document

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    parrafos = [p.text for p in Document(str(salidas["limpia"])).paragraphs]
    cuerpo = [p for p in parrafos if p.startswith("[")]

    con_cargo = [p for p in cuerpo if "CARLA ULLOA — HISTORIADORA" in p]
    solo_nombre = [p for p in cuerpo if "CARLA ULLOA\n" in p and "HISTORIADORA" not in p]
    assert len(con_cargo) == 1
    assert len(solo_nombre) >= 1


def test_la_voz_sin_nombre_no_recibe_uno_inventado(material, tmp_path):
    from docx import Document

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    texto = "\n".join(p.text for p in Document(str(salidas["limpia"])).paragraphs)

    assert "VOZ EN OFF — NARRACIÓN" in texto
    assert "VOZ NO IDENTIFICADA" in texto
    assert "SPEAKER_" not in texto  # las etiquetas técnicas no salen al documento


def test_la_lectura_de_un_texto_se_marca_con_inicio_y_fin(material, tmp_path):
    from docx import Document

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    texto = "\n".join(p.text for p in Document(str(salidas["limpia"])).paragraphs)

    assert "INICIO — LECTURA DE TEXTO" in texto
    assert "FIN — LECTURA DE TEXTO" in texto
    assert "«Confieso que he vivido»" in texto
    assert "Pablo Neruda" in texto


def test_el_xlsx_de_participantes_se_abre_y_trae_la_evidencia(material, tmp_path):
    from openpyxl import load_workbook

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    hoja = load_workbook(salidas["participantes"]).active

    encabezados = [c.value for c in hoja[1]]
    assert "Nivel de confianza" in encabezados and "Evidencia" in encabezados
    filas = list(hoja.iter_rows(min_row=2, values_only=True))
    assert len(filas) == 3
    carla = next(f for f in filas if f[1] == "CARLA ULLOA")
    assert carla[8] == ALTO
    assert "Rótulo en pantalla" in carla[9]
    sin_nombre = next(f for f in filas if f[4] == "SPEAKER_02")
    assert sin_nombre[0] == "VOZ NO IDENTIFICADA"
    assert sin_nombre[9] == "Sin evidencia en el video"


def test_el_xlsx_de_citas_no_atribuye_autor_sin_evidencia(tmp_path, material):
    from openpyxl import load_workbook

    segmentos, identidades, _, contexto = material
    citas = [CitaEnPantalla(inicio_s=10.0, fin_s=14.0, titulo="Sin datos")]
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    fila = next(load_workbook(salidas["citas"]).active.iter_rows(min_row=2, values_only=True))
    assert fila[4] == "No identificado"
    assert fila[7] == "MEDIO"  # sin autor ni año no puede ser ALTO


def test_incertidumbres_lista_lo_que_falta_resolver(material, tmp_path):
    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    texto = salidas["incertidumbres"].read_text(encoding="utf-8")

    assert "Voces sin identificar" in texto
    assert "SPEAKER_02" in texto  # la voz sin nombre se reporta
    assert "Pasajes con reconocimiento dudoso" in texto
    assert "Y eso se nota en la obra" in texto  # el segmento de confianza 0.35
    assert (
        "CARLA ULLOA"
        not in texto.split("Pasajes con reconocimiento")[0].split("Identificaciones que no llegan")[
            0
        ]
    )  # la identificación ALTA no se reporta como problema


def test_el_proceso_tecnico_declara_limites_y_coste(material, tmp_path):
    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas, n_rotulos=7)
    texto = salidas["proceso"].read_text(encoding="utf-8")

    assert "0 USD" in texto
    assert "large-v3-turbo" in texto
    assert "Habla superpuesta" in texto  # el límite real del método
    assert "Voces con nombre propio: 1" in texto
    assert "Voces sin identificar: 2" in texto
    assert "Rótulos leídos en pantalla: 7" in texto


def test_el_srt_tiene_timestamps_validos(material, tmp_path):
    import re

    segmentos, identidades, citas, contexto = material
    salidas = generar_paquete(tmp_path, contexto, segmentos, identidades, citas)
    contenido = salidas["srt"].read_text(encoding="utf-8")

    tiempos = re.findall(r"(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})", contenido)
    assert len(tiempos) == len(construir_intervenciones(segmentos))
    assert tiempos[0][0] == "00:00:00,000"
