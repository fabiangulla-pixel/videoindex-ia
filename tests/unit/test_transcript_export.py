"""Exportación de la transcripción como documento editorial (MD / DOCX / SRT).

Lo que se verifica no es "que escriba un archivo", sino que el documento
diga la verdad: quién habla, en qué minuto, de dónde salió el material y que
es una transcripción automática sin revisar.
"""

from __future__ import annotations

from uuid import uuid4

import pytest

from tests.conftest import hacer_segmentos
from videoindex.application import transcript_export_service as export
from videoindex.domain.models import Video
from videoindex.infrastructure.db.repositories import SegmentRepo, SpeakerRepo, VideoRepo


@pytest.fixture
def video_con_dialogo(con):
    """Entrevista mínima: dos voces alternándose, con procedencia."""
    video = Video(
        video_id=str(uuid4()),
        title="Conversación sobre archivos",
        path="C:/descargas/conversacion.m4a",
        checksum=uuid4().hex,
        duration_seconds=3725.0,  # 1 h 02 min 05 s
        source_url="https://www.youtube.com/watch?v=abc123",
        source_channel="Anales UChile",
        source_published_at="2026-05-20",
    )
    VideoRepo(con).guardar(video)
    segs = hacer_segmentos(
        video.video_id,
        [
            ("¿Cómo empezó el archivo?", 5.0, 9.0),
            ("Empezó en 1998.", 9.0, 12.0),
            ("Con una donación.", 12.0, 15.0),
            ("Qué interesante.", 3600.0, 3602.0),
        ],
    )
    for seg, quien in zip(
        segs, ["SPEAKER_00", "SPEAKER_01", "SPEAKER_01", "SPEAKER_00"], strict=True
    ):
        seg.speaker = quien
    SegmentRepo(con).guardar_lote(segs)
    return video


def test_marca_tiempo():
    assert export.marca_tiempo(0) == "00:00:00"
    assert export.marca_tiempo(3725) == "01:02:05"
    assert export.marca_tiempo(9.25, con_milisegundos=True) == "00:00:09,250"
    assert export.marca_tiempo(-3) == "00:00:00"  # nunca negativo


def test_preparar_agrupa_intervenciones_y_arma_la_ficha(con, video_con_dialogo):
    datos = export.preparar(con, video_con_dialogo.video_id)

    # 4 segmentos → 3 intervenciones (las dos del SPEAKER_01 se unen)
    assert len(datos.intervenciones) == 3
    assert datos.intervenciones[1].texto == "Empezó en 1998. Con una donación."
    assert datos.ficha["Fuente"] == "https://www.youtube.com/watch?v=abc123"
    assert datos.ficha["Canal / autor"] == "Anales UChile"
    assert datos.ficha["Publicado"] == "2026-05-20"
    assert datos.ficha["Duración"] == "01:02:05"


def test_los_nombres_puestos_a_mano_reemplazan_las_etiquetas(con, video_con_dialogo):
    SpeakerRepo(con).renombrar(video_con_dialogo.video_id, "SPEAKER_00", "Entrevistadora")
    datos = export.preparar(con, video_con_dialogo.video_id)
    assert datos.hablantes == ["Entrevistadora", "SPEAKER_01"]

    md = export.a_markdown(datos)
    assert "**Entrevistadora:** ¿Cómo empezó el archivo?" in md
    assert "SPEAKER_01" in md  # el que no se nombró conserva su etiqueta


def test_markdown_lleva_ficha_advertencia_y_marcas_de_tiempo(con, video_con_dialogo):
    md = export.a_markdown(export.preparar(con, video_con_dialogo.video_id))
    assert md.startswith("# Conversación sobre archivos")
    assert export.ADVERTENCIA in md
    assert "`[00:00:05]`" in md
    assert "`[01:00:00]`" in md  # la intervención lejana, en horas


def test_markdown_sin_marcas_de_tiempo(con, video_con_dialogo):
    md = export.a_markdown(export.preparar(con, video_con_dialogo.video_id), False)
    assert "[00:00:05]" not in md
    assert "**SPEAKER_01:**" in md


def test_srt_numera_y_usa_coma_decimal(con, video_con_dialogo):
    srt = export.a_srt(export.preparar(con, video_con_dialogo.video_id))
    lineas = srt.splitlines()
    assert lineas[0] == "1"
    assert lineas[1] == "00:00:05,000 --> 00:00:09,000"
    assert lineas[2].startswith("SPEAKER_00: ")
    assert srt.count("-->") == 3  # un bloque por intervención


def test_video_sin_transcripcion_avisa_en_vez_de_exportar_vacio(con):
    video = Video(video_id=str(uuid4()), title="Sin procesar", path="x.mp4", checksum=uuid4().hex)
    VideoRepo(con).guardar(video)
    with pytest.raises(ValueError, match="todavía no tiene transcripción"):
        export.preparar(con, video.video_id)


def test_video_inexistente(con):
    with pytest.raises(ValueError, match="Video no encontrado"):
        export.preparar(con, "no-existe")


def test_exportar_archivos_reales(con, video_con_dialogo, tmp_path):
    md = export.exportar_markdown(con, video_con_dialogo.video_id, tmp_path / "t.md")
    srt = export.exportar_srt(con, video_con_dialogo.video_id, tmp_path / "t.srt")
    assert md.read_text(encoding="utf-8").startswith("# Conversación")
    assert "-->" in srt.read_text(encoding="utf-8")


def test_docx_se_abre_y_contiene_lo_que_debe(con, video_con_dialogo, tmp_path):
    """No basta con que el archivo exista: se vuelve a abrir y se lee lo que
    quedó dentro, incluida la tabla de procedencia."""
    from docx import Document

    ruta = export.exportar_docx(con, video_con_dialogo.video_id, tmp_path / "t.docx")
    doc = Document(str(ruta))

    parrafos = [p.text for p in doc.paragraphs]
    assert "Conversación sobre archivos" in parrafos[0]
    assert any(export.ADVERTENCIA in p for p in parrafos)
    assert any("SPEAKER_00: ¿Cómo empezó el archivo?" in p for p in parrafos)
    assert any(p.startswith("[00:00:05]") for p in parrafos)

    celdas = [c.text for fila in doc.tables[0].rows for c in fila.cells]
    assert "https://www.youtube.com/watch?v=abc123" in celdas
    assert "Anales UChile" in celdas
